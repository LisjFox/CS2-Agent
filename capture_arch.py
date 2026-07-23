"""截图架构图并嵌入报告"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Paths
arch_html = r'C:\Users\77994\.openclaw\workspace\cs2-agent\architecture.html'
report_path = r'C:\Users\77994\.openclaw\workspace\cs2-agent\课程设计报告_CS2_AI教练.docx'
screenshot_path = r'C:\Users\77994\.openclaw\workspace\cs2-agent\arch_screenshot.png'

# Try to take screenshot with Playwright
try:
    from playwright.sync_api import sync_playwright
    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page(viewport={'width': 1280, 'height': 720})
        page.goto(f'file://{arch_html}')
        page.wait_for_timeout(1000)
        page.screenshot(path=screenshot_path, full_page=True)
        browser.close()
    print(f'Screenshot saved: {screenshot_path}')
    has_screenshot = True
except Exception as e:
    print(f'Screenshot failed: {e}')
    has_screenshot = False

# Generate a matplotlib architecture diagram as backup
try:
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    from matplotlib.patches import FancyBboxPatch
    
    fig, ax = plt.subplots(1, 1, figsize=(10, 7))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 7)
    ax.axis('off')
    fig.patch.set_facecolor('#0f172a')
    ax.set_facecolor('#0f172a')
    
    # Define layers
    layers = [
        (0.5, 5.0, 9.0, 1.8, '#1e3a5f', '用户交互层', 'CLI 对话模式 | HTML 报告 | GSI 实时监控\n@tool 自然语言 → 工具调用'),
        (0.5, 3.0, 9.0, 1.8, '#2e1a47', '智能推理层', 'LangChain 推理引擎 | LLM 教练点评\n残局分析 | 职业对比 | 双轨记忆管理'),
        (0.5, 1.0, 9.0, 1.8, '#1a3a3a', '数据分析层', '击杀热力图 | 武器统计 | 雷达图\n回合分段 | 选手指标计算'),
        (0.5, -1.0, 9.0, 1.8, '#1a2a1a', '数据采集层', '.dem 文件解析 (awpy) | GSI 数据流\n潘一鸣战术知识库 (ChromaDB)'),
    ]
    
    title_color = '#e2e8f0'
    layer_font_color = '#cbd5e1'
    
    for x, y, w, h, color, title, desc in layers:
        rect = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.15",
                              facecolor=color, edgecolor='#475569', linewidth=1.5)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h - 0.35, title, ha='center', va='center',
                fontsize=12, fontweight='bold', color='white',
                fontfamily='sans-serif')
        for li, line in enumerate(desc.split('\n')):
            ax.text(x + w/2, y + h - 0.85 - li*0.35, line, ha='center', va='center',
                    fontsize=9, color='#94a3b8', fontfamily='sans-serif')
    
    # Arrows between layers
    arrow_props = dict(arrowstyle='->', color='#475569', lw=2)
    for mid_y in [4.8, 2.8, 0.8]:
        ax.annotate('', xy=(5, mid_y - 0.1), xytext=(5, mid_y + 0.1),
                    arrowprops=arrow_props)
    
    ax.text(5, 6.8, 'CS2 AI 教练 Agent — 四层架构', ha='center', fontsize=16,
            fontweight='bold', color='white', fontfamily='sans-serif')
    
    plt.tight_layout()
    plt.savefig(screenshot_path, dpi=150, bbox_inches='tight', facecolor='#0f172a')
    plt.close()
    print(f'Architecture diagram saved: {screenshot_path}')
    has_screenshot = True
except Exception as e:
    print(f'Matplotlib diagram failed: {e}')

# Now update the DOCX with the screenshot
if has_screenshot and os.path.exists(screenshot_path):
    doc = Document(report_path)
    
    # Find the architecture section (paragraph 56 is "4.1 架构总览图")
    # Find the right paragraph and insert image after it
    insert_idx = None
    for i, p in enumerate(doc.paragraphs):
        if '4.1 架构总览图' in p.text and 'Heading' in p.style.name:
            insert_idx = i
            break
    
    if insert_idx is not None:
        # Insert image after the heading
        # We need to add it to the body after the heading paragraph
        # In python-docx, we add elements using the paragraph's _element
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        # Add image paragraph
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_para.add_run()
        run.add_picture(screenshot_path, width=Inches(5.5))
        
        # Move the image paragraph after the heading
        heading_elem = doc.paragraphs[insert_idx]._element
        img_elem = img_para._element
        heading_elem.addnext(img_elem)
        
        print(f'Image inserted after paragraph {insert_idx}')
    
    doc.save(report_path)
    print(f'Report updated with screenshot: {report_path}')
else:
    print('No screenshot available, report remains unchanged')
